from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Create new document
doc = Document()

# Set up styles based on SIPEDANG template
styles = doc.styles

# Title style - Cambria 28
title_style = styles['Title']
title_style.font.name = 'Cambria'
title_style.font.size = Pt(28)
title_style.font.bold = False

# Heading 1 - Optima 14
h1_style = styles['Heading 1']
h1_style.font.name = 'Optima'
h1_style.font.size = Pt(14)

# Heading 2 - Optima 11
h2_style = styles['Heading 2']
h2_style.font.name = 'Optima'
h2_style.font.size = Pt(11)

# Normal - Calibri 11
normal_style = styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ==================== CONTENT FROM REV1 ====================

# COVER PAGE
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph("PROPOSAL TEKNIS")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].font.bold = True

subtitle = doc.add_paragraph("PENGADAAN JASA KONSULTANSI PENGEMBANGAN APLIKASI")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph()
main_title = doc.add_paragraph("ONE DATA HUB JHL GROUP")
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
main_title.runs[0].font.size = Pt(24)
main_title.runs[0].font.bold = True

doc.add_paragraph()
year = doc.add_paragraph("TAHUN 2026")
year.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
vendor = doc.add_paragraph("PT Datacaraka Solusindo")
vendor.alignment = WD_ALIGN_PARAGRAPH.CENTER
vendor.runs[0].font.bold = True

address = doc.add_paragraph("Ruko Ketapang Indah Blok B3 No 20\nJl. Kyai H. Zainul Arifin No.20\nJakarta Pusat - Indonesia\nTel: (62 21) 630 1919\nFax: (62 21) 630 5992")
address.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# TABLE OF CONTENTS
toc_title = doc.add_paragraph("Daftar Isi")
toc_title.style = 'Heading 1'

toc_items = [
    ("1. Pendahuluan", "1"),
    ("2. Gambaran Proses Bisnis One Data Hub JHL Group", "2"),
    ("3. Pendekatan Teknis yang Diajukan Aplikasi One Data Hub JHL Group", "3"),
    ("4. Spesifikasi Teknis", "14"),
    ("5. Tahapan Pelaksanaan Pekerjaan", "20"),
]

for item, page in toc_items:
    p = doc.add_paragraph(item)
    p.add_run(f"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t{page}")

doc.add_paragraph()
lof_title = doc.add_paragraph("Daftar Gambar")
lof_title.style = 'Heading 1'

lof_items = [
    "Gambar 1 Proses Bisnis Sistem One Data Hub",
    "Gambar 2 Login One Data Hub",
    "Gambar 3 Form Designer",
    "Gambar 4 Laporan/Reporting",
    "Gambar 5 Halaman Master Data",
    "Gambar 9 Arsitektur Aplikasi",
    "Gambar 11 Arsitektur Sistem",
    "Gambar 15 Timeline Project",
]

for item in lof_items:
    doc.add_paragraph(item)

doc.add_page_break()

# SECTION 1: PENDAHULUAN
doc.add_heading("1. Pendahuluan", level=1)

doc.add_paragraph(
    "PT Datacaraka Solusindo (Winning-soft) merupakan salah satu penyedia jasa Teknologi Informasi "
    "sesuai dengan kebutuhan dari setiap rekan kerja. Kami menyediakan jasa-jasa Teknologi Informasi, "
    "seperti: Sistem Informasi, Pengembangan Website, Jaringan Komputer, serta Multimedia dan Desain Grafis."
)

doc.add_paragraph(
    "One Data Hub merupakan sistem yang dirancang untuk mengonsolidasikan data lintas seluruh unit bisnis, "
    "meningkatkan kualitas data, mempercepat pelaporan, dan menyediakan fondasi analitik (BI), data science, "
    "serta AI/ML yang terukur."
)

doc.add_paragraph(
    "Dalam merespon kebutuhan spesifik dari JHL Group terkait Pengadaan Jasa Konsultansi Pembangunan "
    "One Data Hub, maka PT Datacaraka Solusindo mengusulkan untuk pembuatan sistem web berbasis "
    "pemrograman C# dan web user interface berbasis pemrograman C# .Net Framework dengan basis database "
    "Apache dan tools DevExtreme."
)

doc.add_paragraph("Adapun tujuan dari pekerjaan ini adalah sebagai berikut:")

goals = [
    "Mengonsolidasikan data lintas seluruh unit bisnis",
    "Membantu meningkatkan kualitas data",
    "Mempercepat proses pelaporan data",
    "Membantu proses analisis dalam bentuk Dashboard BI untuk C Level",
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_page_break()

# SECTION 2: GAMBARAN PROSES BISNIS
doc.add_heading("2. Gambaran Proses Bisnis One Data Hub JHL Group", level=1)

doc.add_paragraph(
    "Secara garis besar, proses bisnis yang berjalan untuk sistem procurement adalah sebagai berikut."
)

doc.add_heading("2.1. Gambaran Umum Proses Bisnis Aplikasi One Data Hub JHL Group", level=2)

doc.add_paragraph(
    "Data berasal dari berbagai macam sumber dan dalam berbagai macam format. Kemudian Data akan masuk "
    "ke dalam Data Lakehouse untuk dilakukan pengolahan data dari Raw Data / Bronze -> Clean Data / Silver -> "
    "Fixed Data / Gold. Pada setiap selesai pengolahan data (Bronze and Silver) akan dilakukan rekonsiliasi data, "
    "sampai data yang diolah sudah bersih dan siap digunakan."
)

doc.add_paragraph(
    "Setelah data selesai diolah menjadi Fixed data dalam Data Lakehouse, data tersebut akan masuk ke dalam "
    "Data Warehouse yang kemudian akan disesuaikan untuk Governance, Keamanan, Semantic dan Metric Store-nya."
)

doc.add_paragraph(
    "Setelah diset semuanya, Data siap ditampilkan ke dalam monitoring aplikasi One Data Hub, dan dapat "
    "diakses oleh masing-masing user sesuai dengan hak akses yang sudah diberikan."
)

doc.add_page_break()

# SECTION 3: PENDEKATAN TEKNIS
doc.add_heading("3. Pendekatan Teknis yang Diajukan Aplikasi One Data Hub JHL Group", level=1)

doc.add_heading("3.1 Proses Ingestion Data", level=2)

doc.add_paragraph(
    "Untuk kebutuhan pembuatan form dan komunikasi data antar entitas lebih lancar, maka kami mengusulkan "
    "framework form designer. Form designer akan memiliki fungsi untuk membaca metadata untuk mendapatkan "
    "struktur dan aturan validasi teknis (Referential Integrity) sehingga aplikasi dapat membuat form yang sudah "
    "berisikan validasi seperti wajib diisi, panjang karakter dan tipe datanya."
)

doc.add_paragraph(
    "Selain itu juga form designer dirancang untuk dapat sinkronisasi atau membaca struktur tabel dari repository "
    "yang digunakan sehingga bisa digunakan untuk membuat form lainnya diluar konteks metadata."
)

doc.add_heading("3.1.1 Form Designer dan Form Generator", level=3)

doc.add_paragraph(
    "Form Designer dapat digunakan oleh Admin untuk mengkonfigurasi form inputan. Admin dapat membuat "
    "melakukan sinkronisasi atau generate dengan metadata tersebut sebagai dasarnya. Setelah itu Admin dapat "
    "melakukan setting validasi dan layout webform inputan dengan designer drag and drop."
)

doc.add_paragraph("Form Designer memiliki fitur:")
features = [
    "Designer form dengan drag and drop dan bootstrap layout",
    "Multi Type Field tersedia seperti textfield, textarea, email, numberbox, combobox, currency, WYSIWYG Editor (HTML Editor), Upload File dan lain sebagainya.",
    "Validasi field umum seperti required, value list, pattern, panjang karakter, tipe data, dsb",
    "Form Generator dan List Generator yang dapat digunakan langsung oleh user setelah konfigurasi dipublish (tanpa perlu perubahan code aplikasi).",
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("3.1.2 Validasi Client Side dan Server Side", level=3)
doc.add_paragraph(
    "Validasi dilakukan dua sisi, client side dan server side. Untuk client side menggunakan JavaScript validasi "
    "yang memberikan respons cepat ke user. Untuk server side menggunakan regex validasi pada setiap input "
    "yang masuk ke dalam database."
)

doc.add_heading("3.2 Fitur Utama One Data Hub", level=2)
doc.add_paragraph("Fitur utama yang diusulkan:")
main_features = [
    "Dashboard dengan visualisasi data interaktif",
    "Data ingestion dan ETL pipeline",
    "Data quality management",
    "Data governance dan lineage",
    "API integration dengan sistem existing",
    "Role-based access control",
    "Audit trail lengkap",
    "Export laporan ke Excel dan PDF",
]
for f in main_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# SECTION 4: SPESIFIKASI TEKNIS
doc.add_heading("4. Spesifikasi Teknis", level=1)

doc.add_heading("4.1 Design Arsitektur", level=2)

doc.add_heading("4.1.1 Arsitektur Aplikasi", level=3)
doc.add_paragraph(
    "Arsitektur aplikasi One Data Hub menggunakan pendekatan multilayer architecture yang terdiri dari:"
)
arch_items = [
    "Presentation Layer - Web UI menggunakan C# .Net Framework dengan DevExtreme",
    "Business Logic Layer - Services dan Business Rules menggunakan C#",
    "Data Access Layer - Repository pattern untuk akses database",
    "Integration Layer - API services untuk integrasi external",
]
for item in arch_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("4.1.2 Arsitektur Sistem", level=3)
doc.add_paragraph(
    "Arsitektur sistem mencakup infrastruktur cloud atau on-premise dengan komponen:"
)
sys_items = [
    "Application Server - Windows Server dengan IIS",
    "Database Server - Apache/PostgreSQL",
    "Data Lake - Hadoop atau Cloud Storage",
    "Load Balancer - Untuk high availability",
    "Backup System - Untuk disaster recovery",
]
for item in sys_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("4.2 Konfigurasi Software", level=2)
doc.add_paragraph("Spesifikasi software yang digunakan:")
sw_items = [
    "Backend: C# .Net Framework 8.0",
    "Frontend: C# .Net MVC dengan DevExtreme",
    "Database: Apache/PostgreSQL",
    "IDE: Visual Studio 2022",
    "Version Control: Git",
    "CI/CD: Azure DevOps",
]
for item in sw_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# SECTION 5: TAHAPAN PELAKSANAAN
doc.add_heading("5. Tahapan Pelaksanaan Pekerjaan", level=1)

doc.add_paragraph("Tahapan pelaksanaan pekerjaan meliputi:")

stages = [
    ("Tahap 1 - Analisis dan Desain", [
        "Requirement analysis dan gathering",
        "System design dan architecture",
        "Database design dan ERD",
        "UI/UX design dan wireframing",
    ]),
    ("Tahap 2 - Development", [
        "Setup development environment",
        "Backend development",
        "Frontend development",
        "Integration development",
    ]),
    ("Tahap 3 - Testing", [
        "Unit testing",
        "Integration testing",
        "System testing",
        "User acceptance testing",
    ]),
    ("Tahap 4 - Deployment", [
        "Production deployment",
        "Data migration",
        "User training",
        "Documentation",
    ]),
]

for stage_name, items in stages:
    doc.add_heading(stage_name, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading("5.1 Rencana Implementasi dan Timeline Project", level=2)

# Create timeline table
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
header_cells = table.rows[0].cells
headers = ['Bulan', 'Tahap 1', 'Tahap 2', 'Tahap 3']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
data = [
    ['Bulan 1', 'X', '', ''],
    ['Bulan 2', 'X', 'X', ''],
    ['Bulan 3', '', 'X', ''],
    ['Bulan 4', '', 'X', 'X'],
    ['Bulan 5', '', '', 'X'],
]
for row_idx, row_data in enumerate(data, 1):
    row_cells = table.rows[row_idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row_cells[col_idx].text = cell_data
        for paragraph in row_cells[col_idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_page_break()

# SECTION 6: PENUTUP
doc.add_heading("6. Penutup", level=1)

doc.add_paragraph(
    "PT Datacaraka Solusindo memiliki pengalaman dan track record yang baik dalam pengembangan sistem "
    "informasi enterprise. Dengan tim yang kompeten dan metodologia yang teruji, kami yakin dapat "
    "menyelesaikan proyek One Data Hub JHL Group ini dengan baik."
)

doc.add_paragraph(
    "Kami berharap proposal ini dapat menjadi dasar pertimbangan untuk kerja sama yang sukses "
    "dalam pengembangan sistem One Data Hub JHL Group tahun 2026."
)

doc.add_paragraph()
doc.add_paragraph()

closing = doc.add_paragraph("Jakarta, April 2026")
closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()
doc.add_paragraph()

sign = doc.add_paragraph("PT Datacaraka Solusindo")
sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Save
output_path = "C:\\Users\\budi.kusharyanto\\OneDrive\\WS Dokumen\\Proposal JHL Group-Innovis ref 3.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")