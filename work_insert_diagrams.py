from docx import Document
from docx.shared import Inches
import os, re
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


doc_path = r"C:\Users\budi.kusharyanto\OneDrive\WS Dokumen\Proposal JHL Group-Innovis ref 2.docx"
img_dir = r"C:\Users\budi.kusharyanto\OneDrive\WS Dokumen"

images = [
    ("diagram_gambar_01.png", "Proses Bisnis Sistem One Data Hub", r"Proses Bisnis"),
    ("diagram_gambar_02.png", "Login One Data Hub", r"Login One Data Hub|\\bLogin\\b"),
    ("diagram_gambar_10.png", "Arsitektur Aplikasi", r"Arsitektur Aplikasi"),
    ("diagram_gambar_11.png", "Arsitektur Sistem", r"Arsitektur Sistem"),
    ("diagram_gambar_12.png", "Laporan/Reporting", r"Laporan|Reporting"),
    ("diagram_gambar_13.png", "Halaman Master Data", r"Master Data"),
    ("diagram_gambar_14.png", "Konfigurasi jaringan", r"Konfigurasi jaringan|jaringan"),
    ("diagram_gambar_15.png", "Timeline Project", r"Timeline Project|Timeline"),
    ("diagram_gambar_19.png", "Timeline Project alternative", r"alternative|alternatif"),
]

WIDTH = Inches(5.5)

def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


doc = Document(doc_path)
paras = list(doc.paragraphs)

insertions = []

for img_file, desc, caption_regex in images:
    img_path = os.path.join(img_dir, img_file)
    if not os.path.exists(img_path):
        continue

    m = re.search(r"diagram_gambar_(\d+)", img_file)
    gambar_no = int(m.group(1)) if m else None
    caption_text = f"Gambar {gambar_no}. {desc}" if gambar_no is not None else f"Gambar. {desc}"

    r = re.compile(caption_regex, re.IGNORECASE)
    anchor_idx = None
    for i, p in enumerate(paras):
        if r.search(p.text or ""):
            anchor_idx = i
            break

    if anchor_idx is None:
        continue

    anchor = paras[anchor_idx]

    img_para = insert_paragraph_after(anchor, "")
    run = img_para.add_run()
    run.add_picture(img_path, width=WIDTH)

    insert_paragraph_after(img_para, caption_text)

    insertions.append((img_file, caption_text))

out_path = doc_path.replace(".docx", "_diagrams.docx")
doc.save(out_path)

print(out_path)
for item in insertions:
    print(item[0], '|', item[1])
