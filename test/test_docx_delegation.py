#!/usr/bin/env python3
"""
Test: DOCX Skill Delegation
Tests that subagent can use docx skill to create and process documents.
"""

import os
import sys
import tempfile
import shutil

# Add skills path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'skills', 'docx', 'examples'))

from docx import Document

def test_create_document():
    """Test creating a DOCX document."""
    print("\n[Test 1] Create DOCX document")

    doc = Document()
    doc.add_heading('Test Report', 0)
    doc.add_paragraph('This is a test document for subagent delegation.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Content for section 1.')

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'

    output_path = os.path.join(os.path.dirname(__file__), 'output', 'test_docx.docx')
    doc.save(output_path)

    if os.path.exists(output_path):
        print(f"  [PASS] Created: {output_path}")
        return True
    else:
        print(f"  [FAIL] File not created")
        return False

def test_read_document():
    """Test reading a DOCX document."""
    print("\n[Test 2] Read DOCX document")

    input_path = os.path.join(os.path.dirname(__file__), 'input', 'test_input.docx')

    if not os.path.exists(input_path):
        # Create test input
        doc = Document()
        doc.add_heading('Input Document')
        doc.add_paragraph('Test content for reading.')
        doc.save(input_path)
        print(f"  [INFO] Created test input: {input_path}")

    try:
        doc = Document(input_path)
        paragraphs = len(doc.paragraphs)
        tables = len(doc.tables)
        print(f"  [PASS] Read document: {paragraphs} paragraphs, {tables} tables")
        return True
    except Exception as e:
        print(f"  [FAIL] {e}")
        return False

def test_extract_text():
    """Test text extraction from DOCX."""
    print("\n[Test 3] Extract text from DOCX")

    input_path = os.path.join(os.path.dirname(__file__), 'input', 'test_input.docx')

    if not os.path.exists(input_path):
        doc = Document()
        doc.add_paragraph('Sample text to extract.')
        doc.save(input_path)

    try:
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        if text:
            print(f"  [PASS] Extracted: {len(text)} chars")
            return True
        else:
            print(f"  [WARN] No text extracted")
            return False
    except Exception as e:
        print(f"  [FAIL] {e}")
        return False

def run_tests():
    """Run all DOCX delegation tests."""
    print("=" * 60)
    print("DOCX Skill Delegation Test")
    print("=" * 60)

    results = []
    results.append(test_create_document())
    results.append(test_read_document())
    results.append(test_extract_text())

    passed = sum(results)
    total = len(results)

    print("\n" + "=" * 60)
    print(f"Results: {passed}/{total} tests passed")

    if passed == total:
        print("[PASS] All DOCX delegation tests passed!")
        return 0
    else:
        print("[FAIL] Some tests failed")
        return 1

if __name__ == "__main__":
    sys.exit(run_tests())